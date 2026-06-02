"""Markdown → Word (.docx) converter."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from finance_agent.export.parser import parse_markdown

# 免责声明（追加到报告末尾）
_DISCLAIMER = (
    "免责声明：本报告由 AI 系统基于公开财务数据自动生成，仅供参考，不构成任何投资建议。"
    "报告中的分析和结论基于历史数据和公开市场信息，不保证未来表现。"
    "投资者应结合自身情况独立判断，并咨询专业投资顾问。"
)


def markdown_to_docx(markdown_text: str, output_path: str, stock_name: str = "") -> str:
    """Convert markdown report to Word document.

    Parameters
    ----------
    markdown_text : str
        Full markdown report.
    output_path : str
        Destination file path.
    stock_name : str
        Stock name for the document title.

    Returns
    -------
    str
        The output_path.
    """
    doc = Document()

    # Set default font for the document
    _set_default_font(doc)

    # Title
    title = doc.add_heading(stock_name or "分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.runs[0] if title.runs else None, bold=True, size=18)

    # Body
    sections = parse_markdown(markdown_text)
    for sec in sections:
        if sec.type == "heading":
            # Map markdown heading levels to docx levels (cap at 3)
            level = min(sec.level, 3)
            h = doc.add_heading(sec.text, level=level)
            if h.runs:
                _set_run_font(h.runs[0], bold=True, size=_heading_font_size(level))
        elif sec.type == "table":
            _add_table(doc, sec.rows)
        elif sec.type == "paragraph":
            _add_paragraph(doc, sec.text)
        elif sec.type == "separator":
            doc.add_paragraph()

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(_DISCLAIMER)
    _set_run_font(run, italic=True, size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _set_default_font(doc: Document) -> None:
    """Set default font for the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_run_font(run, bold: bool = False, italic: bool = False, size: int = 10) -> None:
    if run is None:
        return
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _heading_font_size(level: int) -> int:
    return {0: 18, 1: 16, 2: 14, 3: 12}.get(level, 12)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                _set_run_font(run, bold=(i == 0), size=9)


def _add_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with inline bold/italic formatting."""
    p = doc.add_paragraph()
    # Split by **bold** and *italic* patterns
    # Process: **bold**, *italic*, regular text
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            _set_run_font(run, italic=True)
        else:
            run = p.add_run(part)
            _set_run_font(run)
